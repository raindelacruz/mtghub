from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path(__file__).resolve().parents[1] / "deliverables" / "MTGHub_PH_Owner_Operational_Manual.docx"
NAVY = "17324D"
BLUE = "2E74B5"
GOLD = "B8892D"
INK = "24313F"
MUTED = "66717D"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GREEN = "DDEFE4"
AMBER = "FFF3CD"
RED = "F8D7DA"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color=INK, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, 9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.88)
    sec.right_margin = Inches(0.88)
    sec.header_distance = Inches(0.38)
    sec.footer_distance = Inches(0.38)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color, before, after in (
        ("Title", 29, NAVY, 0, 8),
        ("Subtitle", 14, MUTED, 0, 12),
        ("Heading 1", 17, NAVY, 16, 7),
        ("Heading 2", 13.5, BLUE, 11, 5),
        ("Heading 3", 11.5, NAVY, 8, 3),
    ):
        st = styles[name]
        st.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        st._element.rPr.rFonts.set(qn("w:ascii"), st.font.name)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), st.font.name)
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = "Aptos"
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.38)
        st.paragraph_format.first_line_indent = Inches(-0.19)
        st.paragraph_format.space_after = Pt(3.5)
        st.paragraph_format.line_spacing = 1.16


def add_header_footer(section, first=False):
    section.different_first_page_header_footer = first
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("MTGHub PH  |  Owner Operational Manual"), 8.5, bold=True, color=MUTED)
    add_page_field(section.footer.paragraphs[0])


def add_callout(doc, label, text, fill=LIGHT, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(label.upper() + "  "), 9, bold=True, color=color)
    set_run(p.add_run(text), 10.4, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        if isinstance(item, tuple):
            label, text = item
            set_run(p.add_run(label + ": "), 10.5, bold=True, color=NAVY)
            set_run(p.add_run(text), 10.5)
        else:
            set_run(p.add_run(item), 10.5)


def add_steps(doc, steps):
    for title, text in steps:
        p = doc.add_paragraph(style="List Number")
        set_run(p.add_run(title + ". "), 10.5, bold=True, color=NAVY)
        set_run(p.add_run(text), 10.5)


def add_table(doc, headers, rows, widths, header_fill=LIGHT, font_size=9.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], header_fill)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), font_size, bold=True, color=NAVY)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        if r_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(value)), font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_section_title(doc, number, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(f"{number:02d}"), 10, bold=True, color=GOLD)
    h = doc.add_heading(title, level=1)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(10)
        set_run(p2.add_run(subtitle), 10.5, italic=True, color=MUTED)


def add_page_break(doc):
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_header_footer(doc.sections[0], first=True)

    # Cover: editorial-cover pattern with restrained owner-report styling.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("OWNER EDITION"), 10, bold=True, color=GOLD)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("MTGHub PH"), 31, bold=True, color=NAVY, font="Aptos Display")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run("Operational Manual & System Owner's Report"), 17, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("How the marketplace works, who does what, where money moves, and what you must control"), 11.5, italic=True, color=MUTED)
    p.paragraph_format.space_after = Pt(65)
    add_callout(doc, "Purpose", "A practical operating manual based on the implemented repository as reviewed on 19 June 2026. It describes current system behavior, not a proposed future product.", PALE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    set_run(p.add_run("Version 1.0  |  19 June 2026  |  Asia/Manila"), 10, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Confidential - For the Owner and Authorized Administrators"), 9, color=MUTED)

    add_page_break(doc)
    add_section_title(doc, 0, "How to use this manual")
    doc.add_paragraph("This manual is organized around the decisions and routines you need to run MTGHub PH safely. Buyer, seller, and administrator procedures are separated, while shared controls are explained once.")
    add_table(doc, ["Section", "Use it for"], [
        ("Owner brief", "Understand the product, commercial model, and immediate decisions."),
        ("Buyer procedures", "Support users who purchase cards or maintain wanted lists."),
        ("Seller procedures", "Support marketplace sellers and users selling cards to MTGHub."),
        ("Admin procedures", "Moderation, money controls, disputes, data, and platform operations."),
        ("Finance and income", "Understand what is revenue, what is a liability, and what is not monetized."),
        ("Operations and governance", "Run jobs, backups, releases, incidents, retention, and owner reviews."),
    ], [2100, 7260])
    add_callout(doc, "Reading rule", "Where this manual says 'implemented,' the capability is present in the codebase. Actual production health still depends on correct deployment, scheduled jobs, credentials, backups, and operational discipline.", AMBER)
    doc.add_heading("Contents", level=2)
    add_bullets(doc, [
        "1. Owner executive brief", "2. System overview and features", "3. Buyer procedures",
        "4. Seller procedures", "5. Administrator procedures", "6. Money flows and income streams",
        "7. Status lifecycles", "8. Trust, safety, privacy, and control environment",
        "9. Technical operations and business continuity", "10. Owner dashboard and management cadence",
        "11. Priority decisions and known caveats", "12. Quick-reference checklists",
    ])

    add_page_break(doc)
    add_section_title(doc, 1, "Owner executive brief", "The five-minute view of what you own and how it behaves")
    add_callout(doc, "Bottom line", "MTGHub PH is an operational two-sided Magic: The Gathering marketplace plus collection, pricing, wanted-list, and platform buylist tools. It records and controls workflows, but it is not currently an integrated cash-payment processor and it does not automatically earn transaction fees.", LIGHT)
    doc.add_heading("What the system does", level=2)
    add_bullets(doc, [
        ("Marketplace", "Users list cards, buyers cart inventory from one seller, and orders reserve stock transactionally."),
        ("Payment safety", "Buyers submit external payment details and private proof; sellers approve or reject the proof."),
        ("Fulfillment", "Meetup and LBC-style shipping flows include tracking, delivery confirmation, and delayed settlement."),
        ("Store credit", "Wallet balances support checkout, refunds, seller settlement, admin adjustments, and MTGHub buylist payouts."),
        ("Trust", "Reports, disputes, reviews, seller metrics, account moderation, audit logs, and operations events are implemented."),
        ("Card utility", "Card database, Scryfall import/sync, personal collection values, local price history, and wanted-list matching drive engagement."),
    ])
    doc.add_heading("What it does not do", level=2)
    add_bullets(doc, [
        "No automated marketplace commission, listing fee, subscription, advertising fee, or payment-processing fee is implemented.",
        "External cash, GCash, or bank payments are handled outside MTGHub; the system stores references and evidence only.",
        "The platform cannot itself reverse an external payment. Admins record external refund amounts and the responsible party must transfer them outside the platform.",
        "The system does not by itself create an off-site disaster-recovery copy, perform legal review, or provide an independent security assessment.",
    ])
    doc.add_heading("Your owner responsibilities", level=2)
    add_table(doc, ["Responsibility", "Owner outcome"], [
        ("Commercial policy", "Choose monetization, shipping-fee treatment, buylist margins, and store-credit terms."),
        ("Financial control", "Reconcile wallet liabilities, cash payouts, external refunds, inventory, and any logistics collections."),
        ("Trust and safety", "Set response times, evidence standards, escalation authority, and sanctions."),
        ("Production operations", "Ensure jobs, backups, restore tests, monitoring, and release gates actually run."),
        ("Legal and privacy", "Obtain Philippine legal/privacy review and maintain notices, retention, and breach procedures."),
    ], [2500, 6860])

    add_page_break(doc)
    add_section_title(doc, 2, "System overview and features", "A role-based map of the implemented product")
    add_table(doc, ["Capability", "Buyer", "Seller", "Admin / Owner"], [
        ("Accounts", "Register, verify email, profile, reset password", "Same user account", "Roles, status, suspension, audit"),
        ("Cards & prices", "Search cards and view prices", "Use card records for listings", "Create/edit/import/sync cards; add prices"),
        ("Collection", "Track holdings and estimated value", "Can use personally", "No ownership override described"),
        ("Marketplace", "Cart and purchase", "Create and manage listings", "Moderate listings and inspect orders"),
        ("Wanted list", "Publish desired cards; receive offers", "Offer matching active listings", "Review demand"),
        ("Orders", "Pay, upload proof, receive, confirm", "Verify proof, fulfill, track", "Review, dispute support, audit"),
        ("Store credit", "Spend and receive refunds", "Receive settlement", "Credit/debit with notes"),
        ("MTGHub buylist", "Sell cards to the platform", "Same user-side function", "Set offers, inspect, approve, pay"),
        ("Trust", "Message, dispute, review, report", "Message, dispute, receive reviews", "Resolve reports/disputes; moderate reviews"),
    ], [1700, 2500, 2500, 2660], font_size=8.4)
    doc.add_heading("Core operating model", level=2)
    add_steps(doc, [
        ("Attract and organize supply", "Sellers create listings against normalized card records; MTGHub may also publish platform buylist demand."),
        ("Create demand", "Buyers search inventory, track collections and prices, add wanted cards, and receive seller offers."),
        ("Control the transaction", "Checkout reserves inventory, creates an order, debits store credit if used, and records external cash due."),
        ("Verify and fulfill", "Payment proof, messages, tracking, delivery events, and buyer confirmation create an evidence trail."),
        ("Settle and learn", "Wallet settlement is idempotent; reviews, reports, disputes, and metrics feed trust decisions."),
    ])
    add_callout(doc, "Role design", "A normal 'user' can act as both buyer and seller. 'Admin' is the elevated role. Operational separation must therefore come from procedure, access discipline, and audit review—not from separate buyer/seller account types.", AMBER)

    add_page_break(doc)
    add_section_title(doc, 3, "Buyer procedures", "Standard operating procedures for purchasing and wanted-list activity")
    doc.add_heading("3.1 Account readiness", level=2)
    add_steps(doc, [
        ("Register and verify", "Create the account, verify email, and maintain accurate identity and Philippine location details."),
        ("Secure access", "Use a strong password; use the time-limited reset flow if access is lost."),
        ("Review policies", "Read marketplace rules, card condition guide, cancellation/refund policy, privacy notice, and terms before trading."),
    ])
    doc.add_heading("3.2 Purchase from the marketplace", level=2)
    add_steps(doc, [
        ("Select inventory", "Search active listings; verify edition, condition, quantity, location, delivery options, notes, and seller profile."),
        ("Build one-seller cart", "Add the desired quantity. The cart blocks self-purchase, prevents overselling, and permits only one seller per checkout."),
        ("Choose logistics", "Meetup currently adds no logistics fee; LBC adds PHP 100 to the order."),
        ("Choose payment mix", "Apply available store credit up to the allowed amount. The remainder becomes external cash due."),
        ("Place the order", "Checkout locks cart rows and listing quantities, creates the order/items, reduces inventory, debits store credit, and clears the cart in one database transaction."),
        ("Pay on time", "When external cash is due, pay outside MTGHub and submit method, reference, and private payment proof before the 24-hour deadline."),
        ("Monitor proof review", "If rejected, review the seller's notes and resubmit within the renewed payment window. If approved, monitor preparation or meetup readiness."),
        ("Receive and confirm", "For shipping, check tracking and mark delivered; for meetup, confirm receipt. Confirm the order when satisfied."),
        ("Review or dispute", "Leave a verified-purchase review after completion. Open a dispute when evidence supports non-delivery, wrong card, condition mismatch, counterfeit concern, or another material failure."),
    ])
    add_callout(doc, "Buyer warning", "Do not treat payment proof as proof that funds cleared. Buyers should preserve external payment confirmation, messages, packaging, and card photos until the dispute window has passed.", AMBER)
    doc.add_heading("3.3 Wanted list and seller offers", level=2)
    add_steps(doc, [
        ("Add wanted card", "Set desired quantity, optional maximum price, and notes."),
        ("Compare matches", "Review matching active listings and seller offers."),
        ("Accept or decline", "Acceptance adds the offered listing to the cart; it does not bypass normal checkout or payment controls."),
    ])
    doc.add_heading("3.4 Buyer support decision guide", level=2)
    add_table(doc, ["Situation", "Correct response"], [
        ("Order still pending payment", "Buyer may cancel; inventory and used store credit are restored once."),
        ("Deadline passed", "Scheduled expiry should expire the order and restore inventory/credit."),
        ("Payment proof rejected", "Correct proof/reference and resubmit; seller review remains required."),
        ("Item not received or materially wrong", "Open a dispute and preserve evidence; do not make informal wallet edits."),
        ("External refund approved", "Store credit reverses in-system; external cash must be returned outside MTGHub and documented."),
    ], [2600, 6760])

    add_page_break(doc)
    add_section_title(doc, 4, "Seller procedures", "Marketplace selling and selling cards directly to MTGHub")
    doc.add_heading("4.1 Create and maintain a marketplace listing", level=2)
    add_steps(doc, [
        ("Select the exact card", "Use the correct card, set, collector number, language, foil treatment, and condition."),
        ("Describe honestly", "Enter available quantity, PHP unit price, location, delivery options, and defect/condition notes. List only authentic cards you own or are authorized to sell."),
        ("Publish as active", "Only active stock is purchasable and eligible for wanted-list matching."),
        ("Maintain availability", "Edit or cancel promptly when stock changes elsewhere. Do not create misleading duplicate availability."),
    ])
    doc.add_heading("4.2 Process a marketplace sale", level=2)
    add_steps(doc, [
        ("Review the order", "Confirm item snapshots, quantities, logistics, external cash due, and store-credit component."),
        ("Review payment evidence", "For external cash, inspect the private proof and independently confirm funds through the stated channel. Approve or reject with useful notes."),
        ("Prepare fulfillment", "After payment verification, move LBC orders to preparing and then add carrier/reference to mark shipped; move meetup orders to ready for meetup."),
        ("Communicate on-platform", "Keep messages specific to the order and preserve an auditable record."),
        ("Await buyer confirmation", "Buyer marks delivered/confirmed. Delivered orders have a 72-hour automatic confirmation window unless disputed."),
        ("Complete settlement", "Completion requires buyer confirmation. The seller receives only the order's store-credit portion in the wallet; external cash was already paid outside the platform."),
    ])
    add_callout(doc, "Seller control", "Never ship merely because a proof image exists. Confirm cleared external funds first. Proof can be forged or duplicated; the seller is the current payment verifier.", RED, "9B1C1C")
    doc.add_heading("4.3 Send an offer to a wanted-list buyer", level=2)
    add_steps(doc, [
        ("Find a match", "Use an active listing for the exact wanted card."),
        ("Set quantity", "The system caps the offer by listing stock and wanted quantity."),
        ("Wait for acceptance", "Accepted offers enter the buyer's cart and remain subject to normal stock and checkout validation."),
    ])
    doc.add_heading("4.4 Sell cards directly to MTGHub", level=2)
    add_steps(doc, [
        ("Choose a platform buylist entry", "Confirm target card, accepted condition, remaining target quantity, cash offer, and store-credit offer."),
        ("Submit the sell order", "Declare condition, quantity, payout method, and remarks. The order starts as pending receipt."),
        ("Deliver cards", "Use the agreed physical handoff/shipping process and retain proof."),
        ("Wait for inspection", "Admin may accept all, accept part, or reject. Approved quantity and condition determine the final amount."),
        ("Receive payout", "Cash is paid outside the system and marked complete by admin; store credit is posted once to the wallet and ledger."),
    ])

    add_page_break(doc)
    add_section_title(doc, 5, "Administrator procedures", "The operating console for money, trust, content, and production health")
    add_callout(doc, "Admin standard", "Use the UI workflows, require explanatory notes, and preserve the audit trail. Direct database edits to balances, orders, or dispute records should be exceptional, approved, documented, and reconciled.", AMBER)
    doc.add_heading("5.1 Daily opening routine", level=2)
    add_steps(doc, [
        ("Open operations", "Check unresolved critical/error events, pending payments, open disputes, pending deletions, latest backup, and migration state."),
        ("Review transaction queues", "Inspect payment-submitted orders, delayed fulfillment, delivered orders nearing settlement, and MTGHub buylist submissions."),
        ("Review trust queues", "Process reports, disputes, review moderation, and suspicious listings/accounts."),
        ("Check money exceptions", "Investigate negative/odd wallet activity, large manual adjustments, external refunds awaiting confirmation, and buylist cash payouts."),
    ])
    doc.add_heading("5.2 User and access administration", level=2)
    add_bullets(doc, [
        "Promote or demote roles only for a documented operational need; the system blocks removing your own admin role.",
        "Change account status for verified policy or security reasons and record the basis.",
        "Review the immutable admin audit log regularly, especially role, status, wallet, listing, report, dispute, and event actions.",
        "Remove seeded/default credentials before public launch and use named admin accounts rather than shared access.",
    ])
    doc.add_heading("5.3 Catalog and pricing", level=2)
    add_steps(doc, [
        ("Maintain card data", "Create/edit manually or import from Scryfall. Trial bulk sync with a limit before a full sync."),
        ("Maintain price evidence", "Record source, source currency, raw price, PHP conversion, capture date, and notes."),
        ("Review quality", "Correct duplicates, wrong set/collector references, stale prices, and missing images before they affect listings or valuation."),
    ])
    doc.add_heading("5.4 Marketplace moderation", level=2)
    add_steps(doc, [
        ("Assess the evidence", "Review listing content, seller profile, reports, prior orders/reviews, and messages where authorized."),
        ("Apply proportionate status", "Use active, reserved, sold, or cancelled as appropriate; document suspicious-listing actions."),
        ("Protect active transactions", "Do not cancel or alter inventory blindly when an order or dispute is open."),
    ])
    doc.add_heading("5.5 Wallet adjustments", level=2)
    add_steps(doc, [
        ("Verify authority and basis", "Identify the user, direction, amount, business reason, and supporting record."),
        ("Use admin wallet workflow", "Credit or debit with a clear note. The workflow writes a ledger entry and prevents a negative balance."),
        ("Perform second-person review", "For material amounts, another authorized person should compare the ledger, source record, and resulting balance."),
    ])

    add_page_break(doc)
    add_section_title(doc, 5, "Administrator procedures (continued)")
    doc.add_heading("5.6 Dispute resolution", level=2)
    add_steps(doc, [
        ("Freeze and triage", "A dispute freezes automatic settlement. Confirm participants, order stage, stated reason, and requested outcome."),
        ("Gather evidence", "Review listing snapshots, payment proof, payment reference, tracking, order messages, photos/notes, and prior actions."),
        ("Decide consistently", "Choose full refund, partial refund, denied, or no action; write 10-3,000 characters of reasoned notes."),
        ("Allocate refund", "The system restores the approved store-credit portion automatically. Record the external cash portion that must be repaid outside MTGHub."),
        ("Close the loop", "Notify parties, confirm any external transfer evidence, review associated report/review/account action, and retain records."),
    ])
    doc.add_heading("5.7 MTGHub buylist administration", level=2)
    add_steps(doc, [
        ("Publish demand", "Set card, target quantity, accepted condition, cash offer, credit offer, notes, and active state."),
        ("Control receipt", "Mark physical receipt only when custody is confirmed; then move to inspection."),
        ("Inspect line by line", "Record accepted quantity, approved condition, item remarks, and rejected quantity."),
        ("Approve outcome", "Full acceptance, partial acceptance, or rejection determines the approved total using offer snapshots."),
        ("Pay exactly once", "Mark external cash payout complete after transfer, or trigger store-credit payout. Idempotency flags prevent duplicate credit but do not replace reconciliation."),
        ("Record inventory", "The buylist flow increases received quantities; the owner still needs a physical inventory and resale process outside the current marketplace logic."),
    ])
    doc.add_heading("5.8 End-of-day routine", level=2)
    add_bullets(doc, [
        "Confirm no critical operational event remains unexplained.",
        "List unresolved disputes, external refunds, cash buylist payouts, and payment-proof exceptions with an owner and due date.",
        "Reconcile material wallet adjustments and buylist settlements to their source records.",
        "Verify the most recent scheduled backup result and escalate failed jobs immediately.",
    ])

    add_page_break(doc)
    add_section_title(doc, 6, "Money flows and income streams", "The commercial truth of the current implementation")
    add_table(doc, ["Flow", "Who pays whom", "System treatment", "Owner accounting view"], [
        ("External marketplace payment", "Buyer -> seller", "Reference/proof recorded; not processed", "Not platform revenue or platform cash"),
        ("Store-credit checkout", "Buyer wallet -> seller wallet at settlement", "Debited, refundable, then credited once", "Platform liability moves between users"),
        ("LBC logistics fee", "Included in buyer cash due", "Fixed PHP 100 on order", "Policy/accounting destination is not separately defined"),
        ("MTGHub cash buylist", "Platform -> user", "Admin records completion", "Inventory acquisition cost and cash outflow"),
        ("MTGHub credit buylist", "Platform issues credit to user", "Wallet credit posted once", "Inventory acquired; credit liability created"),
        ("Dispute refund", "Seller/responsible party -> buyer", "Credit automated; cash external", "Track outstanding external refund obligation"),
        ("Admin promotion/adjustment", "Platform issues/removes credit", "Ledgered manual entry", "Liability change; expense/revenue depends on basis"),
    ], [1500, 1900, 2900, 3060], font_size=8.0)
    doc.add_heading("Current income stream", level=2)
    add_callout(doc, "Implemented revenue", "There is no automated fee revenue in the marketplace code. The platform buylist can create gross margin only when MTGHub later resells acquired cards above total acquisition, inspection, handling, storage, shrinkage, and selling costs. That resale and cost-accounting process is not fully represented in the current application.", GREEN)
    doc.add_heading("What is not revenue", level=2)
    add_bullets(doc, [
        ("Store-credit balance", "A promise of future value; treat outstanding balances as a liability until used, expired under valid terms, or otherwise settled."),
        ("External seller payment", "Money that never belongs to MTGHub under the current model."),
        ("Refund allocation", "A reversal/obligation, not income."),
        ("Shipping amount", "Do not recognize as platform income until policy identifies the collecting party, actual carrier cost, and any permitted markup."),
    ])
    doc.add_heading("Potential monetization requiring product and policy changes", level=2)
    add_table(doc, ["Option", "Required controls"], [
        ("Transaction commission", "Fee calculation, disclosure, seller net settlement, refunds, tax/accounting, reports, and dispute allocation."),
        ("Listing or subscription plans", "Entitlements, billing, cancellation, invoicing, and account restrictions."),
        ("Payment facilitation", "Regulatory/legal assessment, gateway integration, fees, chargebacks, KYC/AML considerations, reconciliation."),
        ("Logistics margin", "Carrier integration or documented collection/disbursement, service levels, loss claims, and transparent pricing."),
        ("Promoted listings/ads", "Disclosure, ranking policy, billing, content standards, and measurement."),
        ("Buylist resale margin", "SKU inventory, cost basis, grading variance, shrinkage, pricing, sell-through, and gross-margin reporting."),
    ], [2400, 6960])
    add_callout(doc, "Owner decision", "Do not launch a fee merely by changing displayed totals. A legitimate revenue feature needs accounting fields, refund behavior, reports, terms, taxes, and auditable settlement logic.", AMBER)

    add_page_break(doc)
    add_section_title(doc, 7, "Status lifecycles", "The states that determine what users and admins may do")
    doc.add_heading("7.1 Marketplace order lifecycle", level=2)
    add_table(doc, ["Stage", "Status", "Primary actor / next action"], [
        ("Order placed", "pending_payment", "Buyer pays within 24 hours or cancels; job may expire it."),
        ("Evidence sent", "payment_submitted", "Seller reviews private proof."),
        ("Payment accepted", "payment_verified", "Seller starts fulfillment."),
        ("Shipping path", "preparing -> shipped", "Seller adds carrier/tracking; buyer marks delivered."),
        ("Meetup path", "ready_for_meetup", "Buyer confirms delivery/receipt."),
        ("Received", "delivered", "Buyer confirms; otherwise 72-hour auto-confirmation if undisputed."),
        ("Confirmed", "buyer_confirmed", "Seller completes, or scheduled settlement completes automatically."),
        ("Closed", "completed", "Store-credit portion settles once; review becomes available."),
        ("Exception", "cancelled / expired", "Inventory restored; buyer store credit refunded once."),
        ("Trust exception", "disputed / refunded", "Admin resolution controls refund and settlement."),
    ], [1300, 2200, 5860], font_size=8.8)
    doc.add_heading("7.2 Listing lifecycle", level=2)
    add_table(doc, ["Status", "Meaning"], [
        ("active", "Visible and purchasable; eligible for wanted-list matches."),
        ("reserved", "Quantity is committed or unavailable for new purchase."),
        ("sold", "No sellable quantity remains after completed sale."),
        ("cancelled", "Seller/admin removed it; cancelled listings are not restored by order cancellation."),
    ], [1900, 7460])
    doc.add_heading("7.3 Platform buylist lifecycle", level=2)
    add_table(doc, ["Status", "Operational meaning"], [
        ("pending_receipt", "User submitted; MTGHub does not yet confirm custody."),
        ("received", "Physical receipt confirmed."),
        ("under_inspection", "Condition and quantity being assessed."),
        ("accepted / partially_accepted", "Approved total determined; payout pending."),
        ("rejected", "No payout approved."),
        ("completed", "Cash marked paid or store credit posted."),
        ("cancelled", "Submission ended without normal completion."),
    ], [2700, 6660])
    add_callout(doc, "Status discipline", "A status is an accounting and permission control, not just a label. Never force a later state to make a screen look complete; resolve the underlying evidence, inventory, payment, or refund first.", AMBER)

    add_page_break(doc)
    add_section_title(doc, 8, "Trust, safety, privacy, and controls", "What protects the marketplace—and where human judgment remains essential")
    doc.add_heading("Implemented controls", level=2)
    add_bullets(doc, [
        "Password hashing, session authentication, email verification, login throttling, password resets, CSRF checks, and role/participant authorization.",
        "Transactional checkout with row locking, inventory reservation, rollback on failure, and idempotency guards for wallet refunds/settlements.",
        "Private payment-proof storage with MIME/image checks, size/dimension controls, randomized filenames, and participant/admin access restrictions.",
        "Order messages, notifications, status history, reports, disputes, verified-purchase reviews, seller metrics, and review moderation.",
        "Admin audit logs, request-ID JSON logs, database system events, migration checksums, backup checksums, and restore verification tooling.",
        "30-day account-deletion cooling-off and retention cleanup for security logs/payment proofs, subject to disputes or legal hold.",
    ])
    doc.add_heading("Human controls the owner must establish", level=2)
    add_table(doc, ["Control", "Recommended owner rule"], [
        ("Admin access", "Named accounts, least privilege, no shared credentials, prompt removal, periodic access review."),
        ("Wallet adjustments", "Source document + clear note + second review above a defined threshold."),
        ("Disputes", "Written evidence standard, response SLA, conflict-of-interest recusal, consistent remedy matrix."),
        ("Buylist grading", "Condition guide, trained inspector, photo record, variance approval, custody log."),
        ("External cash", "Daily exception report and proof of refunds/payouts; never infer transfer from an admin status."),
        ("Privacy", "Restricted evidence access, retention enforcement, incident contacts, legal review, data-request procedure."),
    ], [2400, 6960])
    doc.add_heading("Retention baseline in current policy/runbook", level=2)
    add_bullets(doc, [
        "Account deletion: 30-day cancellation period; active obligations and non-zero credit must be resolved first.",
        "Security logs: target removal after 90 days.",
        "Closed-order payment proof: target removal after 180 days unless dispute/legal hold applies.",
        "Anonymized transaction, wallet, moderation, dispute, and audit records: may be retained up to seven years.",
    ])
    add_callout(doc, "Legal caveat", "Repository policies are not a substitute for advice on the Philippine Data Privacy Act, consumer protection, taxation, e-commerce, payments, and card authenticity/counterfeit handling.", RED, "9B1C1C")

    add_page_break(doc)
    add_section_title(doc, 9, "Technical operations and business continuity", "The routines that keep the system available, recoverable, and trustworthy")
    doc.add_heading("9.1 Scheduled jobs", level=2)
    add_table(doc, ["Frequency", "Job", "Purpose"], [
        ("Hourly (runbook)", "scripts/expire_orders.php", "Expire overdue payments and settle eligible delivered orders."),
        ("Daily 02:00", "scripts/backup.php", "Database backup plus private-upload manifest and checksum."),
        ("Daily 03:00", "scripts/process_account_deletions.php", "Process eligible deletion requests."),
        ("Daily 03:30", "scripts/cleanup_retention.php", "Apply retention cleanup."),
        ("Weekly after backup", "scripts/verify_restore.php", "Restore to an isolated database and validate it."),
    ], [1900, 3400, 4060], font_size=8.7)
    add_callout(doc, "Documentation mismatch", "README advises running order expiry every five minutes, while the production runbook says hourly. Choose one service level and standardize the scheduler and documentation. Five minutes gives quicker inventory release and settlement; hourly reduces job frequency.", AMBER)
    doc.add_heading("9.2 Safe release procedure", level=2)
    add_steps(doc, [
        ("Enter maintenance mode", "Stop user writes at the web-server layer."),
        ("Back up", "Run the backup and confirm success/checksum."),
        ("Migrate", "Run the versioned migration runner; never edit an already-applied migration."),
        ("Test", "Run production readiness and phase smoke tests in a dedicated staging database."),
        ("Restore service", "Remove maintenance mode and inspect the admin operations dashboard and logs."),
    ])
    doc.add_heading("9.3 Incident response", level=2)
    add_steps(doc, [
        ("Identify", "Use the user-visible request ID, admin operations dashboard, and dated application log."),
        ("Contain", "Freeze affected transactions with a dispute; restrict compromised accounts or credentials."),
        ("Preserve", "Retain logs, proof files, messages, records, and the latest known-good backup."),
        ("Correct", "Use approved workflows; do not repair wallet balances without ledger notes and reconciliation."),
        ("Recover and learn", "Verify service, resolve the system event with documentation, notify affected parties when required, and prevent recurrence."),
    ])
    doc.add_heading("9.4 Backup standard", level=2)
    add_bullets(doc, [
        "Local backups are retained for 30 days, but a local disk alone is not disaster recovery.",
        "Copy encrypted backups to separate restricted storage and test restoration regularly.",
        "Record recovery point objective (maximum acceptable data loss) and recovery time objective (maximum acceptable outage).",
    ])

    add_page_break(doc)
    add_section_title(doc, 10, "Owner dashboard and management cadence", "The measures that tell you whether the marketplace is healthy")
    doc.add_heading("Daily control dashboard", level=2)
    add_table(doc, ["Metric", "Why it matters", "Investigate when"], [
        ("Pending payment / submitted proof", "Cash conversion and inventory lock", "Aged beyond deadline/review SLA"),
        ("Open disputes and reports", "Trust risk and frozen settlement", "Volume or age rises"),
        ("Unresolved critical/error events", "Production health", "Any unexplained critical event"),
        ("Last successful backup", "Recoverability", "Daily backup missing/failed"),
        ("Manual wallet adjustments", "Fraud/error exposure", "High value, repeated, or weak notes"),
        ("Buylist payout queue", "Cash/liability and seller experience", "Accepted but unpaid/uncleared"),
    ], [2200, 3500, 3660], font_size=8.6)
    doc.add_heading("Weekly business review", level=2)
    add_bullets(doc, [
        "Gross merchandise value, completed orders, cancellation/expiry rate, payment-proof approval rate, time to fulfill, dispute rate, and repeat buyer/seller activity.",
        "Active listings, sell-through, wanted-list match/offer acceptance, stale inventory, and seller concentration.",
        "Outstanding store-credit liability, credit issued, credit spent, credit refunded, seller settlement, and reconciliation difference.",
        "Platform buylist: units acquired, approved vs. submitted variance, average cost, cash vs. credit mix, inventory aging, resale proceeds, and realized gross margin.",
    ])
    doc.add_heading("Monthly governance review", level=2)
    add_bullets(doc, [
        "Admin access and audit-log sample; policy consistency; large wallet actions; external refunds/payout evidence.",
        "Backup restoration evidence, job success rate, security events, deletion/retention performance, and capacity trends.",
        "Commercial P&L separated into marketplace activity, platform-owned inventory, logistics, promotions, support, losses, and technology costs.",
    ])
    add_callout(doc, "Owner principle", "Do not manage the business from order totals alone. Separate user-to-user GMV from platform revenue, platform cash, store-credit liability, platform-owned inventory, and external refund obligations.", GREEN)

    add_page_break(doc)
    add_section_title(doc, 11, "Priority decisions and known caveats", "Items requiring owner direction before scale")
    add_table(doc, ["Priority", "Decision / gap", "Owner action"], [
        ("P0", "No implemented monetization", "Select business model; commission and payment handling require designed accounting and policy changes."),
        ("P0", "Default seeded credentials documented", "Change/remove immediately and use named admin accounts."),
        ("P0", "Off-site backup is organizational, not automatic", "Configure encrypted separate storage and assign restore-test ownership."),
        ("P0", "External cash/refunds are outside platform", "Set proof, reconciliation, SLA, and escalation procedures."),
        ("P1", "Expiry schedule conflict: 5 minutes vs hourly", "Choose and standardize one schedule."),
        ("P1", "PHP 100 logistics amount ownership unclear", "Define carrier cost, collector, markup/refund treatment, and reports."),
        ("P1", "Buylist creates inventory but not full inventory accounting", "Implement custody, SKU/cost basis, shrinkage, resale, and margin controls."),
        ("P1", "Store credit is financially material", "Approve terms, liability accounting, promotional budget, limits, and reconciliation."),
        ("P1", "Repository readiness is not independent assurance", "Commission legal/privacy review, penetration test, and capacity test before meaningful public scale."),
        ("P2", "Owner reporting is operational, not full BI", "Create finance and KPI exports/dashboard with period controls."),
    ], [900, 3650, 4810], font_size=8.2)
    doc.add_heading("Current readiness statement", level=2)
    doc.add_paragraph("The repository's Phase 7 readiness report dated 19 June 2026 records passing gates for migrations, authentication/authorization, critical transactions, upload security, backup/restore, policies/lifecycle, and operational monitoring, with zero unresolved critical events after correction. Treat this as repository evidence, not a guarantee that a separate production deployment is configured or healthy.")
    add_callout(doc, "Before public launch", "Run the automated gate in staging, verify production environment variables and SMTP, confirm all scheduled jobs, remove default credentials, complete an off-site restore test, and obtain external legal/security review appropriate to your risk.", RED, "9B1C1C")

    add_page_break(doc)
    add_section_title(doc, 12, "Quick-reference checklists", "Short routines for the moments when speed matters")
    doc.add_heading("Owner launch approval", level=2)
    add_bullets(doc, [
        "Production configuration and public origin verified; secrets are not defaults.",
        "Named admin accounts only; seeded passwords changed/removed.",
        "Migrations and full staging gate pass; no unresolved critical events.",
        "SMTP, payment-proof privacy, upload storage, and access checks verified.",
        "Scheduler jobs confirmed with monitored success/failure notifications.",
        "Latest backup and isolated restore test confirmed; off-site encrypted copy confirmed.",
        "Terms, privacy, marketplace, refund, store-credit, logistics, and monetization policies approved.",
        "Support, dispute, incident, external refund, and buylist payout owners assigned.",
    ])
    doc.add_heading("Admin dispute checklist", level=2)
    add_bullets(doc, [
        "Freeze settlement and confirm participant authorization.", "Preserve listing/order snapshots, payment evidence, tracking, messages, and photos.",
        "Check prior refunds/settlements to prevent duplicate movement.", "Write reasoned resolution notes and calculate credit vs external refund allocation.",
        "Verify external cash refund separately; do not infer it from system status.", "Review related report, review moderation, and account action; close the incident record.",
    ])
    doc.add_heading("Money-control checklist", level=2)
    add_bullets(doc, [
        "Wallet ledger total agrees to wallet balance and source transactions.", "Accepted buylist cash payouts have external transfer proof.",
        "Store-credit buylist payouts and seller settlements occurred once.", "Cancelled/expired/disputed orders restored only the correct inventory and credit.",
        "External refund obligations are aged and followed to confirmation.", "Platform-owned inventory cost and physical count agree.",
    ])
    doc.add_heading("Source basis", level=2)
    doc.add_paragraph("Prepared from the implemented PHP MVC repository, including routes, controllers, models, schema/migrations, policy text, workflow documentation, production runbook, automated test guide, and readiness report. No live production database, bank/payment channel, physical inventory, or scheduler service was independently inspected for this manual.")

    # Ensure all sections receive header/footer and final metadata.
    for section in doc.sections:
        add_header_footer(section, first=(section is doc.sections[0]))
    props = doc.core_properties
    props.title = "MTGHub PH Owner Operational Manual"
    props.subject = "Buyer, seller, administrator, financial, and operational procedures"
    props.author = "MTGHub PH"
    props.keywords = "MTGHub, owner manual, marketplace, operations, buyer, seller, admin"
    props.comments = "Prepared from repository implementation reviewed 19 June 2026."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
